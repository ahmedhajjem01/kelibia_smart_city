"""
Generate Cahier de Charges Word document (GSP-FR-01-01)
for Kelibia Smart City PFE project.
Run with: python generate_cdc.py
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background colour (e.g. 'D0D0D0')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color='000000', sz='12'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, show in sides.items():
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_borders(table, color='000000', sz='12'):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color=color, sz=sz)

def add_run(para, text, bold=False, italic=False, size_pt=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def cell_para(cell, text='', bold=False, italic=False, size_pt=10.5,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Clear cell, add a paragraph with a run, return paragraph."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    if text:
        add_run(p, text, bold=bold, italic=italic, size_pt=size_pt, color=color)
    return p

def set_col_widths(table, widths_mm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Mm(widths_mm[i])

# ─── document setup ───────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Mm(18)
    section.right_margin  = Mm(18)
    section.top_margin    = Mm(15)
    section.bottom_margin = Mm(15)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# ── PAGE 1 ────────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

# ─── HEADER TABLE (logo | title | meta) ───────────────────────────────────────
hdr = doc.add_table(rows=1, cols=3)
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(hdr)
set_col_widths(hdr, [48, 88, 38])

# col 0 — institute name
c0 = hdr.cell(0, 0)
c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c0.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'المعهد العالي للإعلامية بالكاف\n', bold=False, italic=False, size_pt=9)
add_run(p, 'Institut Supérieur\nd\'Informatique du Kef', bold=True, italic=False, size_pt=9.5)

# col 1 — title
c1 = hdr.cell(0, 1)
c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Cahier de charges', bold=True, italic=False, size_pt=18)

# col 2 — ref/date/page (split into 3 sub-rows via nested table)
c2 = hdr.cell(0, 2)
c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# We'll just stack paragraphs
for text, bold in [('GSP-FR-01-01', True), ('Date :\n18/11/2024', False), ('Page 1 sur 2', False)]:
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size_pt=9)
# remove first empty paragraph
c2.paragraphs[0]._element.getparent().remove(c2.paragraphs[0]._element)

doc.add_paragraph()  # small spacer

# ─── SECTION: ÉTUDIANT ────────────────────────────────────────────────────────
def section_header(doc, title):
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t)
    c = t.cell(0, 0)
    set_cell_bg(c, 'D0D0D0')
    p = cell_para(c, title, bold=True, italic=True, size_pt=13,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    return t

section_header(doc, 'Etudiant')

etud = doc.add_table(rows=5, cols=1)
set_table_borders(etud)

# Row 0: Nom étudiant 1 + Monôme/Binôme
p = etud.cell(0, 0).paragraphs[0]
add_run(p, 'Nom & Prénom : ', italic=True, size_pt=11)
add_run(p, 'Hajjem Ahmed', bold=True, size_pt=11)
add_run(p, '          Monôme  □    Binôme  ☑', italic=True, size_pt=11)

# Row 1: Nom étudiant 2
p = etud.cell(1, 0).paragraphs[0]
add_run(p, '                                    ', size_pt=11)
add_run(p, 'Sammoud Maissa', bold=True, size_pt=11)

# Row 2: Spécialité
p = etud.cell(2, 0).paragraphs[0]
add_run(p, 'Spécialité : ', italic=True, size_pt=11)
add_run(p, 'Licence Génie Logiciel', bold=True, size_pt=11)

# Row 3: CIN/Email/Tél étudiant 1
p = etud.cell(3, 0).paragraphs[0]
add_run(p, 'CIN : ', italic=True, size_pt=11)
add_run(p, '14447209', bold=True, size_pt=11)
add_run(p, '    E-mail : ', italic=True, size_pt=11)
add_run(p, 'harounahajjem@gmail.com', bold=True, size_pt=11)
add_run(p, '    Tél. : ', italic=True, size_pt=11)
add_run(p, '56 185 475', bold=True, size_pt=11)

# Row 4: CIN/Email/Tél étudiant 2
p = etud.cell(4, 0).paragraphs[0]
add_run(p, 'CIN : ', italic=True, size_pt=11)
add_run(p, '15350922', bold=True, size_pt=11)
add_run(p, '    E-mail : ', italic=True, size_pt=11)
add_run(p, 'maissasammoud8@gmail.com', bold=True, size_pt=11)
add_run(p, '    Tél. : ', italic=True, size_pt=11)
add_run(p, '20 838 058', bold=True, size_pt=11)

doc.add_paragraph()

# ─── SECTION: SUJET DU PROJET ─────────────────────────────────────────────────
section_header(doc, 'Sujet du projet')

sujet = doc.add_table(rows=1, cols=1)
set_table_borders(sujet)
p = sujet.cell(0, 0).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Kelibia Smart City', bold=True, size_pt=11)
add_run(p, (
    ' — Conception et développement d\'une plateforme numérique bilingue '
    '(Français / Arabe) de services municipaux pour la commune de Kélibia, '
    'Gouvernorat de Nabeul, Tunisie. La plateforme permet aux citoyens d\'accéder '
    'en ligne aux services administratifs (état civil, réclamations géolocalisées, '
    'forum communautaire, actualités municipales) et aux agents municipaux de gérer '
    'ces demandes via un espace dédié, avec un module d\'intelligence artificielle '
    '(NLP) pour la classification automatique des réclamations par catégorie et par '
    'niveau de priorité.'
), size_pt=11)

doc.add_paragraph()

# ─── SECTION: CDC DU PROJET (page 1 part) ─────────────────────────────────────
section_header(doc, 'Cahier de charges du projet')

cdc1 = doc.add_table(rows=1, cols=1)
set_table_borders(cdc1)
c = cdc1.cell(0, 0)

p = c.paragraphs[0]
add_run(p, 'Contexte et définition du problème :', bold=True, size_pt=11.5)

p2 = c.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p2, (
    'La commune de Kélibia, ville côtière du Gouvernorat de Nabeul (Tunisie), gère au quotidien '
    'un volume important de demandes administratives émanant de ses citoyens : certificats de '
    'naissance, attestations de résidence, actes de mariage, déclarations de décès, réclamations '
    'urbaines, etc. Aujourd\'hui, la quasi-totalité de ces démarches s\'effectue encore en '
    'présentiel, générant des files d\'attente, des délais de traitement élevés et une charge de '
    'travail considérable pour les agents municipaux.'
), size_pt=11)

p3 = c.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p3, (
    'L\'absence d\'un outil numérique centralisé empêche tout suivi en temps réel des demandes '
    'citoyennes et limite la transparence et l\'efficacité des services municipaux. Face à cela, '
    'le projet '
), size_pt=11)
add_run(p3, 'Kelibia Smart City', bold=True, italic=True, size_pt=11)
add_run(p3, (
    ' propose une solution web bilingue et responsive, accessible 24h/24, couvrant l\'ensemble '
    'des services municipaux courants et offrant un espace de gestion centralisé aux agents de la '
    'commune, avec un module IA pour l\'automatisation de la classification des réclamations.'
), size_pt=11)

# ══════════════════════════════════════════════════════════════════════════════
# ── PAGE BREAK → PAGE 2 ───────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

# ─── HEADER PAGE 2 ────────────────────────────────────────────────────────────
hdr2 = doc.add_table(rows=1, cols=3)
hdr2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(hdr2)
set_col_widths(hdr2, [48, 88, 38])

c0 = hdr2.cell(0, 0)
c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c0.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'المعهد العالي للإعلامية بالكاف\n', size_pt=9)
add_run(p, 'Institut Supérieur\nd\'Informatique du Kef', bold=True, size_pt=9.5)

c1 = hdr2.cell(0, 1)
c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = c1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Cahier de charges', bold=True, size_pt=18)

c2 = hdr2.cell(0, 2)
c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
for text, bold in [('GSP-FR-01-01', True), ('Date :\n18/11/2024', False), ('Page 2 sur 2', False)]:
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size_pt=9)
c2.paragraphs[0]._element.getparent().remove(c2.paragraphs[0]._element)

doc.add_paragraph()

# ─── SECTION: OBJECTIFS / TECHNIQUES / PLAN ───────────────────────────────────
main = doc.add_table(rows=1, cols=1)
set_table_borders(main)
c = main.cell(0, 0)

# ── Objectifs ──
p = c.paragraphs[0]
add_run(p, 'Objectifs du projet :', bold=True, size_pt=11.5)

p = c.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Objectif principal : ', bold=True, size_pt=11)
add_run(p, (
    'Concevoir et développer une plateforme web complète, bilingue (FR/AR) et responsive, '
    'permettant la dématérialisation des services municipaux de la commune de Kélibia et la '
    'gestion intelligente des réclamations citoyennes.'
), size_pt=11)

p = c.add_paragraph()
add_run(p, 'Objectifs spécifiques :', bold=True, size_pt=11)

objectives = [
    'Permettre aux citoyens de soumettre en ligne des demandes d\'état civil (extrait de naissance, acte de mariage, attestation de résidence, déclaration de décès).',
    'Implémenter un module IA/NLP (TF-IDF + LinearSVC) pour classifier automatiquement les réclamations par catégorie (voirie, éclairage, déchets, bruit…) et par priorité (faible, normale, urgente).',
    'Offrir aux agents municipaux un tableau de bord centralisé pour gérer les demandes, mettre à jour les statuts et consulter les statistiques IA.',
    'Mettre en place un forum citoyen participatif avec votes, notifications et modération agent.',
    'Garantir une interface accessible en français et en arabe (support RTL complet).',
    'Sécuriser la plateforme via authentification JWT, vérification CIN, hachage des mots de passe et CORS.',
    'Déployer la solution sur une infrastructure cloud (Vercel + Neon PostgreSQL) pour une disponibilité 24h/24, 7j/7.',
]
for obj in objectives:
    p = c.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Mm(5)
    add_run(p, obj, size_pt=11)

# ── Techniques ──
p = c.add_paragraph()
add_run(p, 'Techniques (logiciels, matériels, ...) :', bold=True, size_pt=11.5)

techs = [
    ('Backend : ', 'Python 3.12 · Django 5 + Django REST Framework · PostgreSQL · JWT (simplejwt + djoser) · WhiteNoise · Pillow · qrcode'),
    ('Frontend : ', 'React 19 + TypeScript · Vite · React Router v7 · Bootstrap 5 · Leaflet.js (cartographie interactive)'),
    ('Intelligence Artificielle : ', 'scikit-learn (TF-IDF + LinearSVC) · NLTK · NumPy · joblib'),
    ('Infrastructure : ', 'Vercel (serverless, CI/CD GitHub) · Neon PostgreSQL (cloud SSL) · GitHub (versioning)'),
    ('Outils : ', 'VS Code · Postman · pgAdmin · Windows 10/11'),
]
for label, val in techs:
    p = c.add_paragraph()
    p.paragraph_format.left_indent = Mm(3)
    add_run(p, label, bold=True, size_pt=11)
    add_run(p, val, size_pt=11)

# ── Plan de travail ──
p = c.add_paragraph()
add_run(p, 'Plan de travail :', bold=True, size_pt=11.5)

# Plan table inside the cell — we add it to the document then move it
# Easier: just add text rows as paragraphs in a nested table
# Actually python-docx supports nested tables
plan_tbl = OxmlElement('w:tbl')

# We'll build the plan as a simple paragraph-based list for reliability
plan_data = [
    ('1', 'Analyse & Spécification', 'Sem. 1–2', 'Recueil des besoins, rédaction du cahier des charges, validation avec les encadrants'),
    ('2', 'Conception', 'Sem. 3–4', 'Modèle BDD (MCD/MLD), architecture logicielle, maquettes UI, choix technologiques définitifs'),
    ('3', 'Développement Backend', 'Sem. 5–8', 'Django + PostgreSQL, modèles & migrations, API REST (auth, réclamations, état civil, forum, actualités), module IA NLP'),
    ('4', 'Développement Frontend', 'Sem. 9–12', 'React/TypeScript + Vite, pages citoyens & agents, intégration bilingue FR/AR + RTL'),
    ('5', 'Intégration & Tests', 'Sem. 13–14', 'Tests unitaires et fonctionnels, tests d\'intégration, correction bugs, validation utilisateurs'),
    ('6', 'Déploiement & Documentation', 'Sem. 15', 'Déploiement Vercel + Neon DB, rédaction du rapport final, préparation soutenance'),
]

for ph, title, period, tasks in plan_data:
    p = c.add_paragraph()
    p.paragraph_format.left_indent = Mm(3)
    add_run(p, f'Phase {ph} — {title} ({period}) : ', bold=True, size_pt=10.5)
    add_run(p, tasks, size_pt=10.5)

# ─── FOOTER TABLE ─────────────────────────────────────────────────────────────
doc.add_paragraph()

footer = doc.add_table(rows=1, cols=2)
set_table_borders(footer)
set_col_widths(footer, [87, 87])

fl = footer.cell(0, 0)
set_cell_bg(fl, 'F5F5F5')
p = fl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Enseignant chargé du suivi à l\'ISI du Kef', bold=True, italic=True, size_pt=11)

for label, val in [
    ('Nom & Prénom : ', 'Ahmed Toujeni'),
    ('Qualité : ', 'Enseignant / Encadrant PFE'),
    ('E-mail : ', ''),
    ('Signature : ', ''),
]:
    p = fl.add_paragraph()
    add_run(p, label, italic=True, size_pt=11)
    add_run(p, val, bold=True, size_pt=11)

# signature space
for _ in range(3):
    fl.add_paragraph()

fr = footer.cell(0, 1)
set_cell_bg(fr, 'F5F5F5')
p = fr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Réservé à l\'administration', bold=True, italic=True, size_pt=11)
p2 = fr.add_paragraph()
p2.paragraph_format.space_before = Pt(24)
add_run(p2, 'Kef, le ......... / ......... / .........', italic=True, size_pt=11)

# ─── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\MSI\Desktop\kelibia_smart_city\Cahier_des_Charges_Kelibia.docx'
doc.save(out)
print(f'OK Saved: {out}')
